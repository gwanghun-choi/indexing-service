import asyncio
import logging
import os
import re
import shutil
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.parser.base import ParserInterface

logger = logging.getLogger(__name__)


class DocxParser(ParserInterface):
    """
    Word 문서(.docx) 파서

    워드 문서의 텍스트, 이미지, 표를 추출하여 마크다운 형태로 변환합니다.
    """

    def __init__(self):
        """DocxParser 초기화"""
        self.supported_extensions = [".docx"]
        self.temp_dir: Optional[str] = None

    async def parsing(
        self, file_path: str, filename: str = None
    ) -> List[Dict]:
        """
        DOCX 파일을 파싱하여 마크다운 형태로 변환합니다.

        Args:
            file_path: 파일 경로
            filename: 원본 파일 이름 (확장자 확인용)

        Returns:
            List[Dict]: 파싱 결과 목록
        """
        try:
            logger.info(f"DOCX 파일 파싱 시작: file_path={file_path}, filename={filename}")

            # 파일 확장자 확인 - filename에서 먼저 확인
            if filename:
                file_ext = Path(filename).suffix.lower()
            else:
                file_ext = Path(file_path).suffix.lower()

            if not file_ext:
                raise ValueError("파일 확장자를 확인할 수 없습니다")

            if file_ext not in self.supported_extensions:
                raise ValueError(f"지원되지 않는 파일 형식: {file_ext}")

            # 임시 디렉토리 설정
            self.temp_dir = self._setup_temp_directory(Path(file_path).name)

            # 비동기 파싱 실행
            result = await self._parse_docx_async(file_path)

            logger.info(f"DOCX 파일 파싱 완료: {len(result)}개 페이지 처리됨")
            return result

        except Exception as e:
            logger.error(f"DOCX 파싱 중 오류 발생: {e}")
            self._cleanup_temp_directory()
            raise

    async def _parse_docx_async(self, file_path: str) -> List[Dict]:
        """DOCX 파일 비동기 파싱"""
        return await asyncio.to_thread(self._parse_docx_sync, file_path)

    def _parse_docx_sync(self, file_path: str) -> List[Dict]:
        """DOCX 파일 동기 파싱"""
        try:
            # Word 문서 로드 (파일 경로에서 직접 로드)
            doc = Document(file_path)
            self._current_doc = doc  # 하이퍼링크 추출을 위해 저장

            markdown_content = []
            images_info = []

            # 문서의 모든 요소를 순서대로 처리
            for element in doc.element.body:
                if element.tag.endswith("p"):  # 문단
                    for paragraph in doc.paragraphs:
                        if paragraph._element == element:
                            md_text = self._convert_paragraph_to_markdown(paragraph)
                            if md_text.strip():
                                markdown_content.append(md_text)
                            break

                elif element.tag.endswith("tbl"):  # 표
                    for table in doc.tables:
                        if table._element == element:
                            table_md = self._convert_table_to_markdown(table)
                            if table_md.strip():
                                markdown_content.append(table_md)
                            break

            # 이미지 추출
            images_info = self._extract_images(doc, Path(file_path).name)

            # 최종 텍스트 구성
            final_text = "\n\n".join(markdown_content)

            # 특정 패턴 뒤에 구분선 추가
            final_text = self._add_section_dividers(final_text)

            # PDF 파서와 동일한 형식으로 반환
            result = [
                {
                    "page_number": 1,
                    "text": final_text,
                    "images": images_info,
                    "temp_dir": self.temp_dir,
                    "metadata": {
                        "file_name": Path(file_path).name,
                        "total_paragraphs": len(doc.paragraphs),
                        "total_tables": len(doc.tables),
                        "total_images": len(images_info),
                    },
                }
            ]

            return result

        except Exception as e:
            logger.error(f"DOCX 동기 파싱 중 오류: {e}")
            raise

    def _convert_paragraph_to_markdown(self, paragraph: Paragraph) -> str:
        """문단을 마크다운으로 변환"""
        if not paragraph.text.strip():
            return ""

        # Run 레벨에서 포맷팅 적용
        formatted_text = ""
        for run in paragraph.runs:
            text = run.text
            if not text:
                continue

            # 포맷팅 적용
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"

            formatted_text += text

        formatted_text = formatted_text.strip()
        if not formatted_text:
            return ""

        # 연속된 볼드 마크 정리
        formatted_text = self._clean_consecutive_bold_marks(formatted_text)

        # 헤딩 변환 로직을 제거하고 원본 포맷 유지
        # Normal 스타일의 볼드 텍스트는 헤딩이 아니라 일반 텍스트로 처리

        # 리스트 처리
        if formatted_text.startswith("*") and not formatted_text.startswith("**"):
            return f"- {formatted_text[1:].strip()}"

        return formatted_text

    def _clean_consecutive_bold_marks(self, text: str) -> str:
        """연속된 볼드 마크 정리"""
        if not text.strip():
            return ""

        # 1단계: 연속된 볼드 구간들을 하나로 합치기
        # **a****b****c** -> **abc**
        while "****" in text:
            text = re.sub(r"\*\*([^*]*?)\*\*\*\*([^*]*?)\*\*", r"**\1\2**", text)

        # 2단계: ** ** 형태의 빈 볼드나 공백만 있는 볼드 정리
        # ** ** -> 공백, **   ** -> 공백
        text = re.sub(r"\*\*\s*\*\*", " ", text)

        # 3단계: 연속된 공백을 하나로
        text = re.sub(r"\s+", " ", text)

        # 4단계: **텍스트 ** 형태에서 끝의 ** 제거
        text = re.sub(r"(\*\*[^*]+)\s+\*\*", r"\1**", text)

        # 5단계: 마지막 정리 - 앞뒤 불필요한 ** 제거
        text = re.sub(r"^\*\*\s+", "", text)
        text = re.sub(r"\s+\*\*$", "", text)

        # 6단계: ***와 같은 홀수 * 정리
        text = re.sub(r"\*\*\*+", "**", text)

        return text.strip()

    def _add_section_dividers(self, text: str) -> str:
        """특정 패턴 뒤에 구분선 추가"""
        # 패턴 1: "- 신규 DidimAMP 기능 개발" 뒤에 구분선
        text = re.sub(r"(- 신규 DidimAMP 기능 개발)", r"\1\n\n---", text)

        # 패턴 2: "**기타 실적**" 뒤에 구분선
        text = re.sub(r"(\*\*.*?기타 실적.*?\*\*.*)", r"\1\n\n---", text)

        # 패턴 3: 마지막에 구분선 추가
        if not text.endswith("---"):
            text = text.rstrip() + "\n\n---"

        return text

    def _convert_table_to_markdown(self, table: Table) -> str:
        """표를 마크다운 테이블로 변환"""
        if not table.rows:
            return ""

        markdown_rows = []
        all_row_data = []

        # 모든 행 데이터 수집
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_paragraphs = []
                for para in cell.paragraphs:
                    para_text = ""
                    for run in para.runs:
                        text = run.text
                        if run.bold:
                            text = f"**{text}**"
                        elif run.italic:
                            text = f"*{text}*"
                        para_text += text

                    if para_text.strip():
                        # 연속된 볼드 마크 정리
                        para_text = self._clean_consecutive_bold_marks(para_text)

                        # 하위 리스트 처리 (" - 항목" 들여쓰기 보존)
                        if para_text.strip().startswith(" - "):
                            cell_paragraphs.append(para_text.strip())
                        else:
                            cell_paragraphs.append(para_text.strip())

                # 셀 내부 여러 문단을 <br/> 태그로 연결 (마크다운 테이블에서 줄바꿈 표현)
                cell_content = "<br/>".join(cell_paragraphs) if cell_paragraphs else ""

                # 하이퍼링크 추출 및 추가
                hyperlinks = self._extract_hyperlinks_from_cell(cell)
                if hyperlinks:
                    # KMS Wiki 링크가 있으면 추가
                    for link in hyperlinks:
                        if "kms.didimservice.com" in link:
                            cell_content = cell_content.replace(
                                "KMS Wiki 정리 :", f"KMS Wiki 정리 : {link}"
                            )
                            break

                row_cells.append(cell_content)
            all_row_data.append(row_cells)

        if not all_row_data:
            return ""

        # 단일 셀 테이블 처리 (제목이나 리스트)
        if len(all_row_data) == 1 and len(all_row_data[0]) == 1:
            content = all_row_data[0][0]
            if content.startswith("*") and "기능 개발" in content:
                # 리스트로 변환
                items = [item.strip() for item in content.split("*") if item.strip()]
                return "\n".join([f"- {item}" for item in items])
            else:
                return f"\n{content}\n"

        # 일반 테이블 처리
        max_cols = max(len(row) for row in all_row_data)

        # 첫 번째 행을 헤더로 처리
        if all_row_data:
            header_row = all_row_data[0]
            while len(header_row) < max_cols:
                header_row.append("")

            markdown_rows.append("| " + " | ".join(header_row) + " |")
            markdown_rows.append("| " + " | ".join(["---"] * len(header_row)) + " |")

            # 나머지 행들 처리 (병합 셀 고려)
            previous_row = all_row_data[0]  # 헤더 행을 이전 행으로 설정
            for row_data in all_row_data[1:]:
                while len(row_data) < max_cols:
                    row_data.append("")

                # 병합 셀 처리: 이전 행과 같은 값이면 빈 셀로 표시
                processed_row = []
                for col_idx, cell_value in enumerate(row_data):
                    if (
                        col_idx < len(previous_row)
                        and cell_value == previous_row[col_idx]
                        and cell_value.strip()
                    ):
                        processed_row.append("")  # 병합된 셀은 빈 값으로
                    else:
                        processed_row.append(cell_value)

                markdown_rows.append("| " + " | ".join(processed_row) + " |")
                previous_row = row_data  # 현재 행을 다음 비교를 위한 이전 행으로 설정

        return "\n".join(markdown_rows)

    def _extract_hyperlinks_from_cell(self, cell) -> List[str]:
        """셀에서 하이퍼링크 추출"""
        hyperlinks = []
        try:
            for para in cell.paragraphs:
                element = para._element
                # 하이퍼링크 요소 찾기
                for hyperlink in element.iter():
                    if hyperlink.tag.endswith("}hyperlink"):
                        rid = hyperlink.get(qn("r:id"))
                        if rid:
                            try:
                                # 문서의 관계에서 실제 URL 찾기
                                doc = self._get_document_from_element(cell)
                                if doc and rid in doc.part.rels:
                                    rel = doc.part.rels[rid]
                                    hyperlinks.append(rel.target_ref)
                            except Exception:
                                continue
            return hyperlinks
        except Exception:
            return []

    def _get_document_from_element(self, cell):
        """셀에서 문서 객체 가져오기"""
        try:
            # 셀의 부모 테이블을 통해 문서에 접근
            table = cell._element.getparent()
            while table is not None and not table.tag.endswith("}tbl"):
                table = table.getparent()

            if table is not None:
                body = table.getparent()
                while body is not None and not body.tag.endswith("}body"):
                    body = body.getparent()

                if body is not None:
                    # 현재 파싱 중인 문서의 part에서 관계 정보 가져오기
                    return self._current_doc
            return None
        except Exception:
            return None

    def _extract_images(self, doc: DocumentType, file_path: str) -> List[Dict]:
        """이미지 추출"""
        images_info = []

        try:
            rels = doc.part.rels
            image_count = 0

            for _, rel in rels.items():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob

                        # 확장자 결정
                        content_type = rel.target_part.content_type
                        if "png" in content_type:
                            ext = ".png"
                        elif "jpeg" in content_type or "jpg" in content_type:
                            ext = ".jpg"
                        else:
                            ext = ".png"

                        image_count += 1
                        img_filename = f"docx_img_{image_count}{ext}"
                        img_path = os.path.join(self.temp_dir, img_filename)

                        with open(img_path, "wb") as img_file:
                            img_file.write(image_data)

                        # 이미지 크기는 0으로 설정 (실제 크기는 나중에 필요시 확인)
                        width, height = 0, 0

                        images_info.append(
                            {
                                "filename": img_filename,
                                "path": img_path,
                                "width": width,
                                "height": height,
                                "type": "embedded_image",
                                "note": f"Word 문서의 이미지 {image_count}",
                            }
                        )

                    except Exception as e:
                        logger.warning(f"이미지 추출 중 오류: {e}")
                        continue

            return images_info

        except Exception as e:
            logger.error(f"이미지 추출 중 전체 오류: {e}")
            return []

    def _setup_temp_directory(self, file_path: str) -> str:
        """임시 디렉토리 설정"""
        file_name = Path(file_path).stem
        temp_dir = f"temp_docx_images_{file_name}"
        os.makedirs(temp_dir, exist_ok=True)
        return temp_dir

    def _cleanup_temp_directory(self) -> None:
        """임시 디렉토리 정리"""
        if self.temp_dir and os.path.exists(self.temp_dir):
            try:
                shutil.rmtree(self.temp_dir)
                logger.info(f"임시 디렉토리 정리 완료: {self.temp_dir}")
            except Exception as e:
                logger.warning(f"임시 디렉토리 정리 중 오류: {e}")
